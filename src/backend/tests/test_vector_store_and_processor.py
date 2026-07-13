"""Automated tests for Vector Database and Document Processing Pipeline (Checkpoint 7)."""

import os
import shutil
import tempfile
import docx
import fitz  # PyMuPDF
from app.services.document_processor import DocumentProcessor, _evenly_spaced_indices, _is_front_matter
from app.services.vector_store import Document, VectorStore, get_vector_store


def get_auth_headers(client, email: str = "vec_user@example.com"):
    """Register/login user and return auth headers."""
    reg_data = {"email": email, "password": "password123", "full_name": "Vector User"}
    res = client.post("/api/auth/register", json=reg_data)
    if res.status_code == 201:
        res = client.post("/api/auth/verify-email", json={"email": email, "code": "000000"})
    else:
        res = client.post(
            "/api/auth/login", json={"email": email, "password": "password123"}
        )
    token = res.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def test_vector_store_crud_and_search():
    temp_dir = tempfile.mkdtemp()
    try:
        vs = VectorStore(
            collection_name="test_collection_vector", persist_directory=temp_dir
        )
        course_id = "test_vec_1"

        docs = [
            Document(
                content="Python is a powerful programming language used for data science.",
                metadata={
                    "page": 1,
                    "source_file": "intro.pdf",
                    "chunk_id": "c1",
                    "source_chunk_id": f"{course_id}_c1",
                },
            ),
            Document(
                content="FastAPI is a modern web framework for building APIs with Python.",
                metadata={
                    "page": 2,
                    "source_file": "intro.pdf",
                    "chunk_id": "c2",
                    "source_chunk_id": f"{course_id}_c2",
                },
            ),
        ]

        # Add documents
        vs.add_documents(docs, course_id=course_id)
        stats = vs.get_course_stats(course_id)
        assert stats["chunk_count"] == 2
        assert stats["course_id"] == course_id

        # Search for course_id
        res = vs.search("web framework", course_id=course_id, k=5)
        assert len(res) > 0
        assert "FastAPI" in res[0].content
        assert res[0].metadata["course_id"] == course_id
        assert res[0].metadata["page"] == 2
        assert res[0].metadata["source_file"] == "intro.pdf"

        # Search for another course_id (ownership filtering check)
        res_other = vs.search("web framework", course_id="other_course", k=5)
        assert len(res_other) == 0

        # Delete course
        vs.delete_course(course_id)
        stats_after = vs.get_course_stats(course_id)
        assert stats_after["chunk_count"] == 0
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def test_document_processor_txt_pdf_docx():
    temp_dir = tempfile.mkdtemp()
    try:
        vs = VectorStore(
            collection_name="test_collection_proc", persist_directory=temp_dir
        )
        processor = DocumentProcessor(vector_store=vs, chunk_size=50, chunk_overlap=10)

        # 1. Test TXT extraction & chunking
        txt_path = os.path.join(temp_dir, "sample.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(
                "Chapter 1 ................ 5\n\nArtificial Intelligence is transforming education and learning systems globally."
            )

        extracted_txt = processor.extract_text_from_file(txt_path)
        assert len(extracted_txt) == 1
        cleaned_txt = processor.clean_text(extracted_txt[0]["content"])
        assert "................" not in cleaned_txt
        assert "Artificial Intelligence" in cleaned_txt

        chunks_txt = processor.chunk_text(
            cleaned_txt, {"page": 1, "source_file": "sample.txt"}, course_id="course_txt"
        )
        assert len(chunks_txt) >= 1
        assert chunks_txt[0].metadata["source_file"] == "sample.txt"
        assert chunks_txt[0].metadata["course_id"] == "course_txt"
        assert "source_chunk_id" in chunks_txt[0].metadata

        # 2. Test DOCX extraction
        docx_path = os.path.join(temp_dir, "sample.docx")
        doc = docx.Document()
        doc.add_paragraph("Machine learning models require clean training data.")
        doc.save(docx_path)

        extracted_docx = processor.extract_text_from_file(docx_path)
        assert len(extracted_docx) == 1
        assert "Machine learning models" in extracted_docx[0]["content"]

        # 3. Test PDF extraction
        pdf_path = os.path.join(temp_dir, "sample.pdf")
        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        page.insert_text((50, 50), "Deep learning neural networks are powerful.")
        pdf_doc.save(pdf_path)
        pdf_doc.close()

        extracted_pdf = processor.extract_text_from_file(pdf_path)
        assert len(extracted_pdf) == 1
        assert "Deep learning" in extracted_pdf[0]["content"]
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def test_upload_to_processing_pipeline_integration(client):
    headers = get_auth_headers(client, "pipeline@example.com")

    # Create a real text content to upload
    content = (
        "Artificial Intelligence Course Guide.\n"
        "Lesson 1: Introduction to Neural Networks.\n"
        "Neural networks are computing systems inspired by the biological neural networks."
    ).encode("utf-8")

    files = [("files", ("ai_guide.txt", content, "text/plain"))]
    res_upload = client.post("/api/upload", headers=headers, files=files)
    assert res_upload.status_code == 201, res_upload.text
    data = res_upload.json()
    course_id = data["course_id"]

    # In Starlette TestClient, background tasks execute synchronously after upload response
    res_status = client.get(f"/api/course/{course_id}/status", headers=headers)
    assert res_status.status_code == 200, res_status.text
    status_data = res_status.json()

    assert status_data["status"] == "ready"
    assert status_data["stage"] == "completed"
    assert status_data["progress"] == 100
    assert status_data["chunk_count"] > 0
    assert status_data["embedding_status"] == "completed"
    assert status_data["quality_score"] > 0

    # Verify search returns the chunk
    vs = get_vector_store()
    search_res = vs.search("Neural networks", course_id=course_id)
    assert len(search_res) > 0
    assert "Neural networks" in search_res[0].content


def test_delete_course_cleans_chroma(client):
    headers = get_auth_headers(client, "clean_vec@example.com")
    content = b"Some document content to test vector store deletion upon course removal."
    files = [("files", ("delete_vec.txt", content, "text/plain"))]

    res_upload = client.post("/api/upload", headers=headers, files=files)
    course_id = res_upload.json()["course_id"]

    vs = get_vector_store()
    stats_before = vs.get_course_stats(course_id)
    assert stats_before["chunk_count"] > 0

    # Delete course
    res_del = client.delete(f"/api/courses/{course_id}", headers=headers)
    assert res_del.status_code == 200

    # Verify Chroma chunks removed
    stats_after = vs.get_course_stats(course_id)
    assert stats_after["chunk_count"] == 0


def test_chunk_text_merges_content_across_page_boundary():
    """Before cross-page chunking, chunk_text() was called once per page, so two short
    pages could never end up in the same chunk even when well under chunk_size. With
    page_offsets, chunking runs on the whole document at once — same content, fed as one
    string, should merge into a single chunk instead of fragmenting at the page break."""
    processor = DocumentProcessor(vector_store=None, chunk_size=200, chunk_overlap=20)

    page1 = "Xin chao cac ban hoc vien hom nay"
    page2 = "chung ta se hoc ve tri tue nhan tao"
    combined = f"{page1}\n\n{page2}"
    page_offsets = [(0, len(page1), 1), (len(page1) + 2, len(page1) + 2 + len(page2), 2)]

    chunks = processor.chunk_text(
        combined, {"page": 1, "source_file": "doc.pdf"}, course_id="c1", page_offsets=page_offsets
    )

    assert len(chunks) == 1
    assert "hoc vien hom nay" in chunks[0].content
    assert "tri tue nhan tao" in chunks[0].content

    # A chunk starting well inside page 2's offset range must be tagged page 2, not the
    # metadata["page"] fallback (which is page 1 here).
    chunks_p2_start = processor.chunk_text(
        page2, {"page": 1, "source_file": "doc.pdf"}, course_id="c1",
        page_offsets=[(0, len(page2), 2)],
    )
    assert chunks_p2_start[0].metadata["page"] == 2


def test_chunk_text_without_page_offsets_uses_metadata_page():
    """Direct single-page calls (no page_offsets) must behave exactly as before."""
    processor = DocumentProcessor(vector_store=None, chunk_size=200, chunk_overlap=20)
    chunks = processor.chunk_text(
        "Noi dung mot trang duy nhat khong co page_offsets.",
        {"page": 5, "source_file": "doc.pdf"},
        course_id="c1",
    )
    assert len(chunks) == 1
    assert chunks[0].metadata["page"] == 5


def test_strip_repeated_headers_footers():
    """A line repeated identically across >=3 pages of the same document (running
    header/footer) must be stripped; page-specific content must survive untouched."""
    processor = DocumentProcessor(vector_store=None, chunk_size=1800, chunk_overlap=120)
    pages = [
        {"content": "HEADER TITLE\nNoi dung trang 1 rieng biet.\nFOOTER PAGE", "page": 1, "source_file": "doc.pdf"},
        {"content": "HEADER TITLE\nNoi dung trang 2 khac han.\nFOOTER PAGE", "page": 2, "source_file": "doc.pdf"},
        {"content": "HEADER TITLE\nNoi dung trang 3 cung khac.\nFOOTER PAGE", "page": 3, "source_file": "doc.pdf"},
    ]

    result = processor._strip_repeated_headers_footers(pages)

    for original, cleaned in zip(pages, result):
        assert "HEADER TITLE" not in cleaned["content"]
        assert "FOOTER PAGE" not in cleaned["content"]
    assert "Noi dung trang 1" in result[0]["content"]
    assert "Noi dung trang 2" in result[1]["content"]
    assert "Noi dung trang 3" in result[2]["content"]


def test_strip_repeated_headers_footers_leaves_short_page_lists_alone():
    """Fewer than 3 pages isn't enough signal to call a line a repeated header/footer."""
    processor = DocumentProcessor(vector_store=None, chunk_size=1800, chunk_overlap=120)
    pages = [
        {"content": "HEADER TITLE\nNoi dung.", "page": 1, "source_file": "doc.pdf"},
        {"content": "HEADER TITLE\nNoi dung khac.", "page": 2, "source_file": "doc.pdf"},
    ]
    result = processor._strip_repeated_headers_footers(pages)
    assert result == pages


def test_front_matter_classifier_and_ocr_sampling_preserve_later_lesson_pages():
    assert _is_front_matter("Nhà xuất bản Giáo dục. ISBN 978-0-00. Được thẩm định năm 2026.")
    assert _is_front_matter("出版社 审定 版权")
    assert not _is_front_matter("Bài học giải thích cách cộng và trừ các số tự nhiên.")
    assert _evenly_spaced_indices(list(range(30)), 12) == [0, 3, 5, 8, 11, 13, 16, 18, 21, 24, 26, 29]
