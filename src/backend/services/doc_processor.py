"""
Document processing service: Extract text from PDF/DOCX/TXT files.

Optimized for speed:
- Text-layer PDFs skip OCR entirely (fast path).
- Scanned PDFs use OCR at reduced DPI with a page cap.
- Per-step timing is logged for profiling.
"""
import io
import logging
import os
import re
import time
from typing import Any, Callable, Dict, List, Optional

import fitz
import docx
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# OCR is imported lazily to avoid import cost when unused.
_pytesseract = None
_PIL_Image = None

MAX_OCR_PAGES = int(os.getenv("PDF_OCR_MAX_PAGES", "12"))
OCR_DPI = int(os.getenv("PDF_OCR_DPI", "120"))
OCR_ENABLED = os.getenv("PDF_ENABLE_OCR", "true").lower() not in {"0", "false", "no", "off"}
TEXT_MIN_CHARS_PER_PAGE = int(os.getenv("PDF_TEXT_MIN_CHARS_PER_PAGE", "50"))
PDF_SCAN_SAMPLE_PAGES = int(os.getenv("PDF_SCAN_SAMPLE_PAGES", "12"))


def _lazy_ocr_imports():
    """Import pytesseract and PIL only when OCR is actually needed."""
    global _pytesseract, _PIL_Image
    if _pytesseract is None:
        import pytesseract as _pt
        from PIL import Image as _img
        _pytesseract = _pt
        _PIL_Image = _img


def _is_scanned_pdf(doc: fitz.Document) -> bool:
    """Return True if the PDF has no usable text layer at all.

    Scans the first min(20, total) pages; if ANY page has >= 50 chars
    of real text, the PDF is treated as a text-layer PDF.
    """
    sample = min(len(doc), PDF_SCAN_SAMPLE_PAGES)
    for page_num in range(sample):
        raw = doc[page_num].get_text()
        if raw and len(raw.strip()) >= TEXT_MIN_CHARS_PER_PAGE:
            return False
    return True


def _is_corrupted_text(text: str) -> bool:
    """Check if extracted PDF text is corrupted (garbled font encoding or too many replacement characters)."""
    if not text or len(text) < 20:
        return False
    # Count U+FFFD replacement chars and C1/private-use garbage typical of broken
    # font encodings. NOTE: never count(""): that returns len(text)+1 and would
    # flag every page as corrupted.
    replacement_count = text.count("\ufffd") + sum(
        1 for ch in text if "\x80" <= ch <= "\x9f" or "\ue000" <= ch <= "\uf8ff"
    )
    return (replacement_count / len(text)) > 0.05


def _extract_pdf(
    file_path: str,
    on_progress: Optional[Callable[[int, int, str], None]] = None,
) -> List[Document]:
    """Extract text from a PDF file.

    Fast path: read the embedded text layer (instant).
    Slow path: OCR scanned pages at reduced DPI with a page cap.
    """
    doc = fitz.open(file_path)
    total_pages = len(doc)
    logger.info("[PDF] %s – %d pages", file_path, total_pages)

    scanned = _is_scanned_pdf(doc)
    if scanned and OCR_ENABLED:
        logger.info("[PDF] Detected scanned PDF – will OCR up to %d pages at %d DPI", MAX_OCR_PAGES, OCR_DPI)
        _lazy_ocr_imports()
    elif scanned:
        logger.warning("[PDF] Detected scanned PDF, but OCR is disabled. Text extraction may be empty.")
    else:
        logger.info("[PDF] Detected text-layer PDF – skipping OCR")

    documents: List[Document] = []
    t0 = time.perf_counter()
    extracted_chars = 0
    text_layer_pages = 0
    ocr_pages_attempted = 0
    ocr_pages_with_text = 0
    empty_pages = 0
    skipped_ocr_pages = 0

    for page_num in range(total_pages):
        page = doc[page_num]
        p_display = page_num + 1

        # Report progress if callback provided
        if on_progress:
            on_progress(page_num, total_pages, f"Đang đọc trang {p_display}/{total_pages}")

        # --- Fast path: text layer ---
        raw_text = page.get_text()
        text = raw_text.strip() if raw_text else ""

        if text and _is_corrupted_text(text):
            logger.warning(" -> Page %d text layer appears corrupted (encoding errors). Falling back to OCR if available.", p_display)
            if OCR_ENABLED and ocr_pages_attempted < MAX_OCR_PAGES:
                text = ""

        if text:
            source_type = "TEXT_PDF"
            text_layer_pages += 1
        elif scanned and OCR_ENABLED and ocr_pages_attempted < MAX_OCR_PAGES:
            # --- Slow path: OCR fallback for scanned PDFs only ---
            try:
                ocr_pages_attempted += 1
                mat = fitz.Matrix(OCR_DPI / 72, OCR_DPI / 72)
                pix = page.get_pixmap(matrix=mat)
                img = _PIL_Image.open(io.BytesIO(pix.tobytes("png")))
                text = _pytesseract.image_to_string(img, lang="vie+eng", config="--psm 3")
                text = text.strip() if text else ""
                source_type = "IMAGE_PDF"
                if not text:
                    empty_pages += 1
                    continue
                ocr_pages_with_text += 1
            except Exception as e:
                logger.warning(" -> OCR failed on page %d: %s", p_display, e)
                continue
        else:
            # Text-layer PDF with an empty page, or scanned PDF after OCR cap – skip.
            if scanned:
                skipped_ocr_pages += 1
            else:
                empty_pages += 1
            continue

        extracted_chars += len(text)

        content_with_tag = (
            f"=== BẮT ĐẦU DỮ LIỆU TRUY XUẤT ===\n"
            f"[MÃ ĐỊNH DANH TRANG: {p_display}]\n"
            f"NỘI DUNG: {text}\n"
            f"=== KẾT THÚC DỮ LIỆU TRANG {p_display} ==="
        )

        documents.append(Document(
            page_content=content_with_tag,
            metadata={
                "source": file_path,
                "page": p_display,
                "type": source_type,
                "text_chars": len(text),
                "pdf_total_pages": total_pages,
                "scanned_pdf": scanned,
            },
        ))

    doc.close()
    elapsed = time.perf_counter() - t0
    logger.info(
        (
            "[PDF] Extracted %d pages with content in %.2fs "
            "(pages=%d, chars=%d, text_pages=%d, ocr_attempted=%d, ocr_text_pages=%d, empty=%d, skipped_ocr=%d, scanned=%s)"
        ),
        len(documents), elapsed, total_pages, extracted_chars, text_layer_pages,
        ocr_pages_attempted, ocr_pages_with_text, empty_pages, skipped_ocr_pages, scanned,
    )
    return documents


def _extract_docx(file_path: str) -> List[Document]:
    """Extract text from a DOCX file."""
    doc = docx.Document(file_path)
    full_text = [para.text for para in doc.paragraphs]
    text = "\n".join(full_text)
    if not text.strip():
        return []
    return [Document(
        page_content=text,
        metadata={"source": file_path, "type": "DOCX"},
    )]


def _extract_txt(file_path: str) -> List[Document]:
    """Extract text from a plain-text file."""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    if not text.strip():
        return []
    return [Document(
        page_content=text,
        metadata={"source": file_path, "type": "TXT"},
    )]


def get_text_from_any_file(
    file_path: str,
    on_progress: Optional[Callable[[int, int, str], None]] = None,
) -> List[Document]:
    """
    Extract text from PDF (with smart OCR fallback), DOCX, or TXT.
    Returns list of Documents with page metadata.

    Args:
        file_path: Path to the document file.
        on_progress: Optional callback(current, total, message) for progress.
    """
    ext = file_path.rsplit(".", 1)[-1].lower()
    t0 = time.perf_counter()

    try:
        if ext == "pdf":
            documents = _extract_pdf(file_path, on_progress)
        elif ext == "docx":
            documents = _extract_docx(file_path)
        elif ext == "txt":
            documents = _extract_txt(file_path)
        else:
            logger.warning("[DocProcessor] Unsupported extension: %s", ext)
            documents = []
    except Exception as e:
        logger.error("[DocProcessor] Failed to read %s: %s", file_path, e)
        raise

    elapsed = time.perf_counter() - t0
    logger.info("[DocProcessor] %s -> %d documents in %.2fs", file_path, len(documents), elapsed)
    return documents


def analyze_document_quality(
    document_id: str,
    file_path: str,
    documents: List[Document],
) -> Dict[str, Any]:
    """Analyze document extraction quality and classify PDF type."""
    ext = file_path.rsplit(".", 1)[-1].lower()
    page_count = len(documents)
    if ext == "pdf" and documents:
        page_count = int(documents[0].metadata.get("pdf_total_pages", page_count) or page_count)
    if page_count <= 0 and documents:
        page_count = len(documents)

    extracted_char_count = sum(len(doc.page_content) for doc in documents)
    average_chars_per_page = extracted_char_count // max(1, page_count)

    is_scanned_pdf = False
    if ext == "pdf":
        if documents:
            is_scanned_pdf = bool(documents[0].metadata.get("scanned_pdf", False))
        if not is_scanned_pdf and average_chars_per_page < 50:
            is_scanned_pdf = True

    combined_text = "\n".join(doc.page_content for doc in documents)

    # Detect TOC noise (table of contents patterns or dots leaders)
    has_toc_noise = bool(re.search(r"(?:Mục lục|Table of Contents|Contents\b|Chương\s+\d+[\s\.]{4,}\d+)", combined_text, re.IGNORECASE))
    has_dot_leaders = bool(re.search(r"\.{4,}|\.\s+\.\s+\.\s+\.", combined_text))
    has_many_page_numbers = len(re.findall(r"^\s*\d+\s*$", combined_text, re.MULTILINE)) >= max(3, page_count // 2)
    has_broken_spacing = bool(re.search(r"\b(?:vềPython|dữliệu|trítuệ|sửdụng|khôngcó|ngônngữ|dựán|biến thểchính)\b", combined_text))
    has_code = bool(re.search(r"\b(?:def |class |#include|import |public class |function |return |const |let |var )\b", combined_text))
    has_formulas = bool(re.search(r"(\b[EFPQW]\s*\([\w,\s]+\)|\\\w+|[≤≥≠×Σ∫√±])", combined_text))
    has_tables = combined_text.count("|") >= 6 or bool(re.search(r"(?:\t.{1,30}){3,}", combined_text))

    # Language detection
    vi_diacritics = len(re.findall(r"[àáãạảăắằẳẵặâấầẩẫậèéẹẻẽêềếểễệđìíĩỉịòóõọỏôốồổỗộơớờởỡợùúũụủưứừửữựỳýỵỷỹ]", combined_text, re.IGNORECASE))
    en_words = len(re.findall(r"\b(?:the|and|for|with|that|have|this|from|which|would|there|their|what|about|when)\b", combined_text, re.IGNORECASE))
    if vi_diacritics > 20 and en_words > 20:
        detected_language = "mixed"
    elif vi_diacritics > 15:
        detected_language = "vi"
    elif en_words > 15 or (extracted_char_count > 100 and vi_diacritics == 0):
        detected_language = "en"
    else:
        detected_language = "unknown"

    # Quality score computation
    quality_score = 100
    warnings: List[str] = []

    if not documents or extracted_char_count == 0:
        quality_score = 0
        warnings.append("Tài liệu trống hoặc không thể trích xuất văn bản.")
    else:
        if average_chars_per_page < 50:
            quality_score -= 55
            warnings.append("Lượng văn bản trên trang rất thấp (<50 ký tự/trang), có thể là tài liệu scan hoặc ảnh.")
        elif average_chars_per_page < 150:
            quality_score -= 20
            warnings.append("Lượng văn bản trên trang khá thấp.")

        if is_scanned_pdf:
            quality_score -= 25
            if "Lượng văn bản trên trang rất thấp (<50 ký tự/trang), có thể là tài liệu scan hoặc ảnh." not in warnings:
                warnings.append("Phát hiện tài liệu PDF dạng scan hoặc hình ảnh.")

        if has_toc_noise or has_dot_leaders:
            quality_score -= 10
            warnings.append("Phát hiện nhiễu định dạng mục lục hoặc dòng chấm lửng rải rác.")

        if has_broken_spacing:
            quality_score -= 10
            warnings.append("Phát hiện lỗi dính từ hoặc khoảng trắng bị ngắt quãng.")

        if has_many_page_numbers:
            quality_score -= 5

    quality_score = max(0, min(100, quality_score))

    # Clean context score: how structurally clean the extracted text is (noise-focused).
    clean_context_score = 100
    if has_toc_noise or has_dot_leaders:
        clean_context_score -= 25
    if has_broken_spacing:
        clean_context_score -= 20
    if has_many_page_numbers:
        clean_context_score -= 15
    if is_scanned_pdf:
        clean_context_score -= 20
    clean_context_score = max(0, min(100, clean_context_score))

    # Content depth score: how much substantive content is available.
    if not documents or extracted_char_count == 0:
        content_depth_score = 0
    else:
        content_depth_score = min(100, int((extracted_char_count / 6000) * 60) + min(40, page_count * 2))
    content_depth_score = max(0, min(100, content_depth_score))

    # PDF Type classification
    if ext != "pdf":
        pdf_type = "mixed_content"
    elif not documents or quality_score <= 20 or (is_scanned_pdf and average_chars_per_page < 50):
        pdf_type = "poor_quality_pdf"
    elif is_scanned_pdf:
        pdf_type = "scanned_pdf"
    elif has_code:
        pdf_type = "code_document"
    elif average_chars_per_page < 450 and (re.search(r"^[•\-\*\+]\s+", combined_text, re.MULTILINE) or page_count >= 10):
        pdf_type = "lecture_slides"
    elif re.search(r"\b(?:Abstract|Tóm tắt|References|Tài liệu tham khảo|Keywords|Phương pháp nghiên cứu)\b", combined_text, re.IGNORECASE):
        pdf_type = "academic_paper"
    elif average_chars_per_page >= 800 or page_count >= 15:
        pdf_type = "text_book"
    else:
        pdf_type = "mixed_content"

    document_type = {
        "text_book": "textbook",
        "lecture_slides": "lecture_slides",
        "academic_paper": "academic_paper",
        "code_document": "code_document",
        "scanned_pdf": "scanned_pdf",
        "poor_quality_pdf": "poor_quality",
        "mixed_content": "mixed_content",
    }.get(pdf_type, "unknown")

    # Recommended action
    if not documents or extracted_char_count < 200 or quality_score < 40:
        if is_scanned_pdf:
            recommended_action = "needs_ocr"
        else:
            recommended_action = "insufficient_context"
    elif quality_score < 70 or warnings:
        recommended_action = "generate_with_warning"
    else:
        recommended_action = "generate"

    # Per-output generation readiness enum
    if is_scanned_pdf and average_chars_per_page < 50 and extracted_char_count < 500:
        generation_readiness = "needs_ocr"
    elif not documents or extracted_char_count == 0:
        generation_readiness = "insufficient"
    elif quality_score >= 75 and extracted_char_count >= 1500 and page_count >= 3:
        generation_readiness = "full"
    elif quality_score >= 50 and extracted_char_count >= 800:
        generation_readiness = "limited"
    elif extracted_char_count >= 500:
        generation_readiness = "summary_only"
    else:
        generation_readiness = "insufficient"

    recommended_actions: List[str] = []
    if generation_readiness == "needs_ocr":
        recommended_actions.append("Bật OCR hoặc sử dụng bản PDF có lớp văn bản (text layer) rõ ràng hơn.")
    elif generation_readiness == "insufficient":
        recommended_actions.append("Tải lại tài liệu có nhiều nội dung văn bản hơn hoặc kiểm tra file gốc.")
    elif generation_readiness == "summary_only":
        recommended_actions.append("Sử dụng các tính năng tạo tóm tắt, dàn ý (outline) hoặc học trọng tâm để khai thác dữ liệu hiện có.")
    elif generation_readiness == "limited":
        recommended_actions.append("Có thể tạo học liệu rút gọn; cân nhắc bổ sung tài liệu để có bản đầy đủ hơn.")
    else:
        recommended_actions.append("Có thể tạo đầy đủ học liệu (Sách, Slides, Quiz, Video) từ tài liệu này.")
    if has_toc_noise or has_dot_leaders:
        recommended_actions.append("Hệ thống sẽ tự động loại bỏ các đoạn mục lục/chấm lửng khỏi ngữ cảnh sinh nội dung.")
    if has_broken_spacing:
        recommended_actions.append("Hệ thống sẽ tự động chuẩn hóa lỗi dính từ tiếng Việt trong quá trình làm sạch.")

    return {
        "document_id": document_id,
        "file_name": os.path.basename(file_path),
        "file_type": ext,
        "page_count": page_count,
        "text_extraction_success": len(documents) > 0 and extracted_char_count > 0,
        "extracted_char_count": extracted_char_count,
        "average_chars_per_page": average_chars_per_page,
        "detected_language": detected_language,
        "document_type": document_type,
        "is_scanned_pdf": is_scanned_pdf,
        "is_scanned_or_image_based": is_scanned_pdf,
        "has_toc_noise": has_toc_noise,
        "has_dot_leaders": has_dot_leaders,
        "has_many_page_numbers": has_many_page_numbers,
        "has_broken_spacing": has_broken_spacing,
        "has_code": has_code,
        "has_formulas": has_formulas,
        "has_tables": has_tables,
        "quality_score": quality_score,
        "clean_context_score": clean_context_score,
        "content_depth_score": content_depth_score,
        "generation_readiness": generation_readiness,
        "warnings": warnings,
        "recommended_action": recommended_action,
        "recommended_actions": recommended_actions,
        "pdf_type": pdf_type,
    }

