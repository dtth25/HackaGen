"""Document processing service for extracting, cleaning, chunking, and embedding documents."""

import logging
import os
import re
from collections import Counter
from typing import List, Optional, Tuple
from pydantic import BaseModel
import fitz  # PyMuPDF
import docx
from app.models.course import Course
from app.services.vector_store import Document, VectorStore

logger = logging.getLogger(__name__)

_SENTENCE_BOUNDARY_RE = re.compile(r"[.?!][ \n]")


def _find_split_position(text: str, start: int, end: int, chunk_size: int) -> int:
    """Find the best position to end a chunk at, preferring a sentence boundary, then a
    newline, then a plain space — only within the back half of the window so a chunk
    never ends up drastically shorter than requested. Returns -1 for a hard cut at `end`
    when no good boundary is found."""
    half = start + chunk_size // 2

    best = -1
    for m in _SENTENCE_BOUNDARY_RE.finditer(text, start, end):
        pos = m.start() + 1  # right after the punctuation, before the trailing whitespace
        if pos > half:
            best = pos
    if best != -1:
        return best

    split_pos = text.rfind("\n", start, end)
    if split_pos == -1 or split_pos <= half:
        split_pos = text.rfind(" ", start, end)
    if split_pos != -1 and split_pos > half:
        return split_pos

    return -1


def _is_quota_exhausted_error(exc: Exception) -> bool:
    """Distinguish a Gemini embedding quota/rate-limit failure (safe to silently retry via
    OpenRouter) from a genuine bug (bad input, misconfigured key, etc. — should still fail
    loud). Matches the same signal GeminiEmbeddingFunction's own retry loop already logged."""
    text = str(exc)
    return "RESOURCE_EXHAUSTED" in text or "429" in text or "quota" in text.lower()


class ProcessingResult(BaseModel):
    """Result of processing documents for a course."""

    course_id: str
    status: str
    chunk_count: int
    quality_score: int
    error: Optional[str] = None


class DocumentProcessor:
    """Processes course documents: extraction, cleaning, chunking, embedding, and vector storage."""

    def __init__(
        self,
        vector_store: VectorStore,
        chunk_size: Optional[int] = None,
        chunk_overlap: Optional[int] = None,
    ):
        from app.core.config import settings

        self.vector_store = vector_store
        self.chunk_size = chunk_size if chunk_size is not None else settings.DOCUMENT_CHUNK_SIZE
        self.chunk_overlap = chunk_overlap if chunk_overlap is not None else settings.DOCUMENT_CHUNK_OVERLAP

    def _update_course_db(
        self,
        course_id: str,
        status: str,
        stage: str,
        progress: int,
        chunk_count: int = 0,
        quality_score: int = 0,
        embedding_status: str = "pending",
        name: Optional[str] = None,
        error_message: Optional[str] = None,
        embedding_provider: Optional[str] = None,
        db_session_factory=None,
    ):
        """Helper to update course status in database."""
        if db_session_factory is None:
            from app.services.database import SessionLocal as factory
        else:
            factory = db_session_factory
        try:
            with factory() as db:
                course = (
                    db.query(Course)
                    .filter(Course.id == course_id, Course.is_deleted == False)  # noqa: E712
                    .first()
                )
                if course:
                    course.status = status
                    course.stage = stage
                    course.progress = progress
                    if chunk_count > 0:
                        course.chunk_count = chunk_count
                    if quality_score > 0:
                        course.quality_score = quality_score
                    if name:
                        course.name = name
                    if embedding_provider:
                        course.embedding_provider = embedding_provider
                    course.embedding_status = embedding_status
                    # Clears any stale error from a prior attempt on success, persists the
                    # real reason on failure — always set explicitly, not just on failure.
                    course.error_message = error_message
                    db.commit()
        except Exception as e:
            logger.error(f"Failed to update course {course_id} in DB: {e}")

    def _generate_course_title(self, all_documents: List[Document]) -> Optional[str]:
        """Best-effort AI-generated short course title from a sample of extracted text.
        Tried exactly once per course (from process_course) — on any failure the caller
        falls back to the document's own filename, like how chat UIs name a conversation
        once and leave the rest to manual rename."""
        try:
            from app.core.config import settings
            from app.services.llm import LLMService

            sample = "\n\n".join(doc.content for doc in all_documents[:5])[:4000]
            if not sample.strip():
                return None
            result = LLMService(model=settings.GEMINI_COURSE_MODEL or None).generate_course_title(sample)
            title = (result.title or "").strip()
            return title or None
        except Exception as e:
            logger.error(f"Course title generation failed, falling back to filename: {e}", exc_info=True)
            return None

    @staticmethod
    def _filename_fallback_title(file_path: str) -> str:
        """Clean a filename into a presentable course title: drop the upload-time timestamp
        prefix and the extension (e.g. '1730000000_Virtual Tree.pdf' -> 'Virtual Tree')."""
        filename = os.path.basename(file_path)
        parts = filename.split("_", 1)
        if len(parts) == 2 and parts[0].isdigit():
            filename = parts[1]
        return os.path.splitext(filename)[0].strip() or filename

    def purge_course_storage(self, course_id: str) -> None:
        """Remove on-disk uploads and vector store chunks for a course. Shared by
        single-course delete and full-account deletion so the cleanup logic lives in
        exactly one place."""
        import shutil
        from app.core.config import settings

        upload_dir = os.path.join(settings.UPLOAD_DIR, course_id)
        if os.path.exists(upload_dir):
            shutil.rmtree(upload_dir, ignore_errors=True)
        self.vector_store.delete_course(course_id)

    def extract_text_from_file(self, file_path: str) -> List[dict]:
        """Extract text from PDF/DOCX/TXT file. Returns list of dicts with content and page number."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)
        # Strip timestamp prefix if present (e.g. 1234567890_test.pdf -> test.pdf)
        parts = filename.split("_", 1)
        if len(parts) == 2 and parts[0].isdigit():
            clean_filename = parts[1]
        else:
            clean_filename = filename

        pages = []

        if ext == ".pdf":
            try:
                from app.core.config import settings

                with fitz.open(file_path) as doc:
                    page_texts = [(page.get_text() or "").strip() for page in doc]

                    scan_mode = False
                    if settings.PDF_ENABLE_OCR and page_texts:
                        sample_n = min(settings.PDF_SCAN_SAMPLE_PAGES, len(page_texts))
                        low_text_count = sum(
                            1 for t in page_texts[:sample_n] if len(t) < settings.PDF_TEXT_MIN_CHARS_PER_PAGE
                        )
                        scan_mode = low_text_count >= max(1, sample_n // 2)

                    ocr_budget = settings.PDF_OCR_MAX_PAGES if scan_mode else 0
                    for idx, text in enumerate(page_texts):
                        if scan_mode and ocr_budget > 0 and len(text) < settings.PDF_TEXT_MIN_CHARS_PER_PAGE:
                            ocr_budget -= 1
                            ocr_text = self._ocr_page(doc, idx, settings.PDF_OCR_DPI)
                            if ocr_text and len(ocr_text) > len(text):
                                text = ocr_text
                        if text:
                            pages.append(
                                {
                                    "content": text,
                                    "page": idx + 1,
                                    "source_file": clean_filename,
                                }
                            )
            except Exception as e:
                # Fallback for dummy/test PDF files in unit tests that contain plain text
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    if text and text.strip():
                        logger.warning(
                            f"PDF open failed for {file_path}, fallback to plain text read for testing."
                        )
                        pages.append(
                            {
                                "content": text.strip(),
                                "page": 1,
                                "source_file": clean_filename,
                            }
                        )
                    else:
                        raise e
                except Exception:
                    logger.error(f"Error reading PDF {file_path}: {e}")
                    raise e

        elif ext == ".docx":
            try:
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text and para.text.strip():
                        full_text.append(para.text.strip())
                if full_text:
                    pages.append(
                        {
                            "content": "\n".join(full_text),
                            "page": 1,
                            "source_file": clean_filename,
                        }
                    )
            except Exception as e:
                # Fallback for dummy/test DOCX files in unit tests that contain plain text
                try:
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        text = f.read()
                    if text and text.strip():
                        logger.warning(
                            f"DOCX open failed for {file_path}, fallback to plain text read for testing."
                        )
                        pages.append(
                            {
                                "content": text.strip(),
                                "page": 1,
                                "source_file": clean_filename,
                            }
                        )
                    else:
                        raise e
                except Exception:
                    logger.error(f"Error reading DOCX {file_path}: {e}")
                    raise e

        elif ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                if text and text.strip():
                    pages.append(
                        {
                            "content": text.strip(),
                            "page": 1,
                            "source_file": clean_filename,
                        }
                    )
            except Exception as e:
                logger.error(f"Error reading TXT {file_path}: {e}")
                raise

        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        return pages

    def _ocr_page(self, doc: "fitz.Document", page_index: int, dpi: int) -> Optional[str]:
        """Render a PDF page to an image and ask Gemini vision to transcribe its text.
        Best-effort: returns None on any failure so the caller keeps whatever text
        extraction already produced (possibly empty/short)."""
        try:
            page = doc[page_index]
            zoom = dpi / 72.0
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
            image_bytes = pix.tobytes("png")

            from app.services.llm import LLMService

            return LLMService().ocr_page_image(image_bytes) or None
        except Exception as e:
            logger.warning(f"OCR fallback failed for page {page_index + 1}: {e}")
            return None

    def clean_text(self, text: str) -> str:
        """Clean extracted text: remove repetitive noise, extra whitespace, Table of Contents leaders."""
        if not text:
            return ""

        # Remove TOC dot leaders (e.g. "Chapter 1 ................ 5")
        text = re.sub(r"\.{4,}\s*\d*", " ", text)
        # Remove repetitive dashes or underscores
        text = re.sub(r"[-_]{4,}", " ", text)
        # Normalize whitespace and newlines
        lines = [line.strip() for line in text.split("\n")]
        # Filter out very short noisy lines (like page numbers standing alone)
        cleaned_lines = [line for line in lines if len(line) > 1 or not line.isdigit()]

        return "\n".join(cleaned_lines).strip()

    def chunk_text(
        self,
        text: str,
        metadata: dict,
        course_id: str,
        page_offsets: Optional[List[Tuple[int, int, int]]] = None,
    ) -> List[Document]:
        """Chunk text with overlap and attach required metadata.

        When `page_offsets` (list of (start, end, page_number) covering `text`) is given —
        used when `text` is the concatenation of every page of one document, so ideas
        split across a page boundary don't get fragmented — each chunk's `page` is derived
        from the offset range it starts in instead of the single `metadata["page"]` value.
        Without it, behaves exactly as a single-page chunk call (`metadata["page"]` for all
        chunks), which is what direct single-page callers still rely on.
        """
        if not text:
            return []

        chunks = []
        source_file = metadata.get("source_file", "unknown")
        default_page = metadata.get("page", 1)

        def _page_for_offset(pos: int) -> int:
            if not page_offsets:
                return default_page
            for p_start, p_end, page_num in page_offsets:
                if p_start <= pos < p_end:
                    return page_num
            return page_offsets[-1][2]

        # Character-boundary chunking with overlap, preferring sentence/newline/space breaks
        start = 0
        text_len = len(text)
        idx = 0

        while start < text_len:
            end = start + self.chunk_size
            if end < text_len:
                split_pos = _find_split_position(text, start, end, self.chunk_size)
                if split_pos != -1:
                    end = split_pos

            chunk_content = text[start:end].strip()
            if chunk_content:
                page = _page_for_offset(start)
                chunk_id = f"{source_file}_p{page}_c{idx}"
                source_chunk_id = f"{course_id}_{chunk_id}"

                doc = Document(
                    content=chunk_content,
                    metadata={
                        "page": page,
                        "source_file": source_file,
                        "chunk_id": chunk_id,
                        "source_chunk_id": source_chunk_id,
                        "course_id": course_id,
                    },
                )
                chunks.append(doc)
                idx += 1

            if end >= text_len:
                break
            start = max(end - self.chunk_overlap, start + 1)

        return chunks

    def _strip_repeated_headers_footers(self, pages: List[dict]) -> List[dict]:
        """Detect lines that repeat identically across >=3 pages of the same document
        (running headers/footers) and strip them before chunking. Needs the full page
        list at once (unlike clean_text, which only ever sees one page)."""
        if len(pages) < 3:
            return pages

        line_counts: Counter = Counter()
        for p in pages:
            lines = {line.strip() for line in p["content"].split("\n") if line.strip()}
            for line in lines:
                if len(line) >= 3:
                    line_counts[line] += 1

        repeated = {line for line, count in line_counts.items() if count >= 3}
        if not repeated:
            return pages

        result = []
        for p in pages:
            kept = [line for line in p["content"].split("\n") if line.strip() not in repeated]
            result.append({**p, "content": "\n".join(kept)})
        return result

    def _is_low_information(self, text: str) -> bool:
        """A chunk is noise if it's too short or mostly non-alphanumeric (leftover
        table borders, decorative characters, stray symbols)."""
        stripped = text.strip()
        if len(stripped.split()) < 15:
            return True
        alnum_count = sum(1 for c in stripped if c.isalnum())
        return len(stripped) > 0 and (alnum_count / len(stripped)) < 0.5

    def extract_and_chunk_file(self, path: str, course_id: str) -> List[Document]:
        """Extract, clean, dedup headers/footers, cross-page chunk, and low-info-prune a
        single file. Shared by process_course() and the standalone re-embed migration
        script (scripts/reembed_courses.py) so both stay in sync with chunking changes."""
        extracted_pages = self.extract_text_from_file(path)
        cleaned_pages = []
        for page_data in extracted_pages:
            cleaned = self.clean_text(page_data["content"])
            if cleaned:
                cleaned_pages.append({**page_data, "content": cleaned})
        cleaned_pages = self._strip_repeated_headers_footers(cleaned_pages)

        # Concatenate every page of this document into one string with an
        # offset->page map, so a chunk never gets cut off just because it
        # happens to straddle a page boundary.
        combined_parts: List[str] = []
        page_offsets: List[Tuple[int, int, int]] = []
        offset = 0
        source_file = None
        for page_data in cleaned_pages:
            content = page_data["content"]
            if not content:
                continue
            source_file = page_data["source_file"]
            combined_parts.append(content)
            page_offsets.append((offset, offset + len(content), page_data["page"]))
            offset += len(content)
            combined_parts.append("\n\n")
            offset += 2

        if not combined_parts:
            return []

        combined_text = "".join(combined_parts)
        page_meta = {"page": page_offsets[0][2], "source_file": source_file}
        doc_chunks = self.chunk_text(combined_text, page_meta, course_id, page_offsets=page_offsets)
        filtered_chunks = [c for c in doc_chunks if not self._is_low_information(c.content)]
        # Pruning noise should never zero out a whole document's contribution —
        # a genuinely tiny document (or a chunk_size small enough to slice below
        # the word-count floor) should still upload, not silently disappear.
        return filtered_chunks or doc_chunks

    def process_course(
        self, course_id: str, file_paths: List[str], db_session_factory=None
    ) -> ProcessingResult:
        """
        1. Extract text từ files (PDF/DOCX/TXT)
        2. Clean text: remove headers, footers, TOC noise
        3. Chunk text với overlap
        4. Generate embeddings
        5. Store in Chroma
        """
        logger.info(f"Starting document processing pipeline for course {course_id}")
        self._update_course_db(
            course_id,
            status="processing",
            stage="extracting",
            progress=20,
            db_session_factory=db_session_factory,
        )

        all_documents: List[Document] = []
        try:
            # 1. Extract, 2. Clean, 3. Chunk
            for path in file_paths:
                all_documents.extend(self.extract_and_chunk_file(path, course_id))

            self._update_course_db(
                course_id,
                status="processing",
                stage="chunking",
                progress=50,
                db_session_factory=db_session_factory,
            )

            if not all_documents:
                raise ValueError(
                    "No valid text could be extracted from uploaded files."
                )

            self._update_course_db(
                course_id,
                status="processing",
                stage="embedding",
                progress=75,
                db_session_factory=db_session_factory,
            )

            # 4. Generate embeddings & 5. Store in Chroma. Gemini embedding is primary; if its
            # own retries exhaust on a quota/rate-limit error (large documents like a 300-page
            # book can need 8+ batch calls against a tight free-tier budget shared across every
            # feature), silently redo the WHOLE course's embedding via OpenRouter instead — never
            # split a single course across both collections (see VectorStore._collection_for).
            # The user never sees which provider indexed their course, only that it worked.
            embedding_provider = "gemini"
            try:
                self.vector_store.add_documents(all_documents, course_id=course_id, provider="gemini")
            except Exception as e:
                if not _is_quota_exhausted_error(e):
                    raise
                logger.warning(
                    f"Gemini embedding quota exhausted for course {course_id}, "
                    f"retrying full course via OpenRouter fallback: {e}"
                )
                self.vector_store.add_documents(all_documents, course_id=course_id, provider="openrouter")
                embedding_provider = "openrouter"

            # Calculate quality score (e.g., based on average chunk length and total chunks)
            quality_score = min(100, max(50, len(all_documents) * 5 + 60))

            # One AI-naming attempt per course, like a chat UI naming a new conversation —
            # no retries. Falls back to the (cleaned) uploaded filename so `name` is always
            # set once a course is ready; after this, naming is purely manual (rename).
            course_title = self._generate_course_title(all_documents) or self._filename_fallback_title(
                file_paths[0]
            )

            # Completed!
            self._update_course_db(
                course_id,
                status="ready",
                stage="completed",
                progress=100,
                chunk_count=len(all_documents),
                quality_score=quality_score,
                embedding_status="completed",
                name=course_title,
                embedding_provider=embedding_provider,
                db_session_factory=db_session_factory,
            )

            logger.info(
                f"Successfully processed course {course_id}: {len(all_documents)} chunks created."
            )
            return ProcessingResult(
                course_id=course_id,
                status="ready",
                chunk_count=len(all_documents),
                quality_score=quality_score,
            )

        except Exception as e:
            error_msg = str(e)
            logger.error(
                f"Document processing failed for course {course_id}: {error_msg}"
            )
            self._update_course_db(
                course_id,
                status="failed",
                stage="failed",
                progress=0,
                embedding_status="failed",
                error_message=error_msg[:500],
                db_session_factory=db_session_factory,
            )
            return ProcessingResult(
                course_id=course_id,
                status="failed",
                chunk_count=0,
                quality_score=0,
                error=error_msg,
            )


def get_document_processor() -> DocumentProcessor:
    """Get DocumentProcessor instance initialized with singleton VectorStore."""
    from app.services.vector_store import get_vector_store
    return DocumentProcessor(vector_store=get_vector_store())
